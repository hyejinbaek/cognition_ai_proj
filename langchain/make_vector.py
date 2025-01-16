# vector store에 vector값 생성
import os
import json
import uuid
from dotenv import load_dotenv
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_core.prompts import PromptTemplate
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain.document_loaders import PyMuPDFLoader, TextLoader
from flask import Flask, request, jsonify, session
import unicodedata
from langchain.schema import Document
from docx import Document as DocxDocument
from datetime import datetime
import pandas as pd

openai_api_key = os.getenv("OPENAI_API_KEY")

fine_tuned_model = 'ft:gpt-4o-2024-08-06:auton::ASh95CCN'


# 엑셀 파일 로드 함수
def load_excel(file_path):
    print(f"Loading Excel file: {file_path}")  # 디버깅: 엑셀 파일 경로 출력
    try:
        # pandas로 엑셀 파일을 로드하여 데이터프레임으로 변환
        df = pd.read_excel(file_path, engine='openpyxl')  # 엑셀 파일 읽기
        print(f"Excel file loaded successfully: {file_path}")  # 디버깅: 로드 성공 출력
        # 각 행을 문서로 변환하여 페이지 콘텐츠에 추가
        full_text = []
        for index, row in df.iterrows():
            full_text.append(" ".join(str(cell) for cell in row))  # 각 셀의 값을 공백으로 구분하여 결합
        return [Document(page_content="\n".join(full_text))]
    except Exception as e:
        print(f"Error loading Excel file: {e}")  # 디버깅: 오류 출력
        return []
    
# .docx 파일 로드
def load_docx(file_path):
    doc = DocxDocument(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return [Document(page_content="\n".join(full_text))]

def load_document(file_path):
    file_path = unicodedata.normalize('NFKC', file_path)
    print(f"Checking file type: {file_path}")  # 디버깅: 파일 유형 확인
    if file_path.endswith('.pdf'):
        loader = PyMuPDFLoader(file_path=file_path)
        documents = loader.load()
    elif file_path.endswith('.txt'):
        loader = TextLoader(file_path=file_path)
        documents = loader.load()
    elif file_path.endswith('.docx'):
        documents = load_docx(file_path)
    elif file_path.endswith('.xlsx'):
        documents = load_excel(file_path)
    else:
        raise ValueError("지원되지 않는 파일 형식입니다.")
    return documents

loaded_files = {}  # 파일 수정 시간 기록을 위한 딕셔너리

def load_all_documents_in_folder(folder_path):
    docs = []
    print(f"Listing files in folder: {folder_path}")  # 디버깅: 폴더 내 파일 목록 출력
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if filename.endswith(('.pdf', '.txt', '.docx', '.xlsx')):  # 엑셀 파일도 포함
            try:
                # 새 파일이나 수정된 파일만 로드
                print(f"Loading file: {file_path}")  # 디버깅: 파일 경로 출력
                if file_path not in loaded_files or os.path.getmtime(file_path) > loaded_files[file_path]:
                    docs.extend(load_document(file_path))
                    loaded_files[file_path] = os.path.getmtime(file_path)
            except Exception as e:
                print(f"{filename} 로드 중 오류: {e}")
    return docs

# 로컬에서 데이터 처리
def create_vectorstore_local(folder_path, output_path):
    docs = load_all_documents_in_folder(folder_path)

    if not docs:
        print("로드된 문서가 없습니다. 문제를 확인해 주세요.")
        return
    
    # 문서 분할
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200, length_function=len)
    split_documents = text_splitter.split_documents(docs)
    print(f"문서 분할 완료: {len(split_documents)} 개 문서")

    # 임베딩 및 벡터 스토어 생성
    embeddings = OpenAIEmbeddings()
    vectorstore = FAISS.from_documents(documents=split_documents, embedding=embeddings)

    # 벡터 스토어 저장
    vectorstore.save_local(output_path)
    print(f"벡터 스토어가 {output_path} 경로에 저장되었습니다.")

if __name__ == '__main__':
    # 환경 변수 로드
    load_dotenv()

    # 문서를 로드할 폴더 경로
    folder_path = "../d"  # 로드할 파일들이 저장된 폴더 경로
    output_path = "./vectorstore"  # 벡터 스토어를 저장할 경로

    # 폴더 경로와 벡터 스토어 저장 경로가 올바른지 확인
    if not os.path.exists(folder_path):
        print(f"지정된 폴더가 존재하지 않습니다: {folder_path}")
    else:
        # 벡터 스토어 생성 및 저장
        print("문서를 로드하고 벡터 스토어를 생성합니다...")
        create_vectorstore_local(folder_path, output_path)
        print("프로세스가 완료되었습니다.")