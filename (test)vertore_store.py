# 폴더 전체 읽어와서 langchain

import os
from dotenv import load_dotenv
import pandas as pd

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_core.prompts import PromptTemplate
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain.document_loaders import PyMuPDFLoader, PyPDFLoader, TextLoader
from docx import Document
from langchain.docstore.document import Document

load_dotenv()

openai_api_key = os.getenv("OPENAI_API_KEY")

fine_tuned_model = 'ft:gpt-4o-2024-08-06:auton::ASh95CCN'

# 다양한 파일 유형을 처리하는 함수 정의
def load_document(file_path):
    if file_path.endswith('.pdf'):
        loader = PyMuPDFLoader(file_path=file_path)
        return loader.load()
    elif file_path.endswith('.txt'):
        loader = TextLoader(file_path=file_path)
        return loader.load()
    elif file_path.endswith('.docx'):
        return load_docx(file_path)
    elif file_path.endswith('.xlsx'):
        return load_xlsx(file_path)
    else:
        raise ValueError("지원되지 않는 파일 형식입니다.")

# .docx 파일을 로드하는 함수 정의
def load_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return [{"text": "\n".join(full_text)}]


def load_xlsx(file_path):
    df = pd.read_excel(file_path, engine='openpyxl')
    # 각 행을 문서로 변환
    documents = []
    for index, row in df.iterrows():
        row_text = " ".join([f"{col}: {row[col]}" for col in df.columns if pd.notna(row[col])])
        documents.append({"text": row_text})
    return documents

# 폴더 내 모든 파일을 로드하는 함수 정의
def load_all_documents_in_folder(folder_path):
    docs = []
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        try:
            docs.extend(load_document(file_path))
            print(f"{filename} 파일이 성공적으로 로드되었습니다.")
        except Exception as e:
            print(f"{filename} 파일을 로드하는 중 오류 발생: {e}")
    return docs

# 폴더 경로 지정
folder_path = "./dataset"

# 폴더 내 모든 파일을 로드하여 docs 리스트에 저장
docs = load_all_documents_in_folder(folder_path)
print(f"총 문서 수: {len(docs)}")

# 사전을 LangChain의 Document 객체로 변환
langchain_docs = [
    Document(page_content=doc['text'], metadata={}) for doc in docs
]

# 단계 2: 문서 분할(Split Documents)
text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=50)
split_documents = text_splitter.split_documents(langchain_docs)  # langchain_docs 사용

print(f"총 분할된 문서 수: {len(split_documents)}")

# 단계 3: 임베딩(Embedding) 생성
embeddings = OpenAIEmbeddings()

# 단계 4: DB 생성(Create DB) 및 저장
# 벡터스토어를 생성합니다.
vectorstore = FAISS.from_documents(documents=split_documents, embedding=embeddings)

# 단계 5: 검색기(Retriever) 생성
# 문서에 포함되어 있는 정보를 검색하고 생성합니다.
retriever = vectorstore.as_retriever()

# 검색기에 쿼리를 날려 검색된 chunk 결과를 확인합니다.
retriever.invoke("요청이 (비업무)개인시간_흡연 등, 요청 사유가 담배일 경우는?")

# 단계 6: 프롬프트 생성(Create Prompt)
# 프롬프트를 생성합니다.
prompt = PromptTemplate.from_template(
    """
이것은 이석사유에 대해 승인/거절/보류를 판단하는 AI모델이다.
요청 카테고리에는 (비업무)개인시간_흡연 등, (업무)회의, (업무)기타업무, (업무)출장,이동,외근 이렇게 4가지가 있다.
요청 카테고리와 요청 사유가 일치해야 승인이 된다.
만약 일치하지 않다면 거절로 판단하고 거절사유도 함께 명시해야한다.
답을 모르면 보류로 판단하면 된다.

#Question: 
{request_type, request_reason} 
#Context: 
{response} 

#Answer:"""
)

# 단계 7: 언어모델(LLM) 생성
# 모델(LLM) 을 생성합니다.
llm = ChatOpenAI(model_name=fine_tuned_model, temperature=0)

# 단계 8: 체인(Chain) 생성
chain = (
    {"context": retriever, "question": RunnablePassthrough()}
    | prompt
    | llm
    | StrOutputParser()
)

# 사용자 입력 받기 및 체인 실행
def ask_question(request_type, request_reason):
    request_type = chain.invoke(request_type)
    request_reason = chain.invoke(request_reason)
    response = chain.invoke(request_type, request_reason)
    print(f"질문: {request_type, request_reason}")
    print(f"답변: {response}")

if __name__ == '__main__':
    # 예시 질문
    ask_question('(비업무)개인시간_흡연 등', '담배')